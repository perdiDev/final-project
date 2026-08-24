#!/usr/bin/env python3
"""Build a single academic .docx combining finished skripsi chapters (BAB I-IV)."""
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BUILD_DIR = os.path.dirname(os.path.abspath(__file__))
DRAFT_DIR = os.path.dirname(BUILD_DIR)
OUT_PATH = os.path.join(DRAFT_DIR, "Skripsi-Gabungan-BAB-I-IV.docx")
DIAGRAMS_DIR = os.path.join(BUILD_DIR, "diagrams")

CHAPTER_FILES = [
    f"{BUILD_DIR}/BAB-1.clean.md",
    f"{BUILD_DIR}/BAB-2.clean.md",
    f"{BUILD_DIR}/BAB-3.clean.md",
    f"{BUILD_DIR}/BAB-4.clean.md",
]

TOKEN_RE = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)")
IMAGE_LINE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)$")

# Mermaid diagrams cannot be rendered natively by python-docx; each was hand-translated
# to Graphviz DOT (see diagrams/*.dot) and pre-rendered to PNG (diagrams/*.png) because
# the sandbox has no network access for mermaid-cli/puppeteer. Match by a unique
# substring of the mermaid source so the right image is picked regardless of which
# chapter/position it appears in.
MERMAID_IMAGE_MAP = [
    ("Kamera ZED", "pipeline_deepstream.png",
     "Diagram alur pipeline DeepStream inti beserta thread pencatatan benchmark/deteksi dan parser tegrastats."),
    ("NMS bawaan nvinfer", "nms_position.png",
     "Perbandingan posisi tahap NMS pada varian baseline versus EfficientNMS_TRT di dalam pipeline."),
    ("Build-time - sekali, offline", "nms_buildtime_runtime.png",
     "Alur fusi EfficientNMS_TRT pada tahap build-time (offline) dan eksekusinya pada tahap runtime (per frame)."),
    ("Desain eksperimen", "measurement_structure.png",
     "Struktur pengukuran: dari desain eksperimen menuju kelompok metrik runtime dan rumusan masalah yang dijawab."),
]


# ---------- Markdown -> block parsing ----------

def parse_blocks(text):
    blocks = []
    current_para = []
    current_list = None  # [type, [items]]
    current_table = []
    in_fence = False
    fence_lang = ""
    fence_lines = []

    def flush_para():
        nonlocal current_para
        if current_para:
            blocks.append(("para", " ".join(current_para)))
            current_para = []

    def flush_list():
        nonlocal current_list
        if current_list:
            blocks.append(("list", current_list))
            current_list = None

    def flush_table():
        nonlocal current_table
        if current_table:
            blocks.append(("table", current_table[:]))
            current_table = []

    def flush_all():
        flush_para()
        flush_list()
        flush_table()

    for raw in text.split("\n"):
        line = raw.rstrip("\n")

        if in_fence:
            if line.strip().startswith("```"):
                in_fence = False
                if fence_lang.strip().lower() == "mermaid":
                    blocks.append(("mermaid", "\n".join(fence_lines)))
                else:
                    blocks.append(("code", "\n".join(fence_lines)))
                fence_lines = []
                fence_lang = ""
            else:
                fence_lines.append(line)
            continue

        if line.strip() == "":
            flush_all()
            continue
        if line.strip().startswith("```"):
            flush_all()
            in_fence = True
            fence_lang = line.strip()[3:]
            fence_lines = []
            continue
        img_match = IMAGE_LINE_RE.match(line.strip())
        if img_match:
            flush_all()
            blocks.append(("image", (img_match.group(1), img_match.group(2))))
            continue
        if line.startswith("# "):
            flush_all()
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            flush_all()
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            flush_all()
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("|"):
            flush_para()
            flush_list()
            current_table.append(line.strip())
        elif re.match(r"^\d+\.\s", line):
            flush_para()
            flush_table()
            if current_list is None or current_list[0] != "num":
                flush_list()
                current_list = ["num", []]
            current_list[1].append(re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("- "):
            flush_para()
            flush_table()
            if current_list is None or current_list[0] != "bul":
                flush_list()
                current_list = ["bul", []]
            current_list[1].append(line[2:])
        elif line.startswith(" ") and current_list is not None:
            current_list[1][-1] += " " + line.strip()
        else:
            if current_table:
                flush_table()
            if current_list:
                flush_list()
            current_para.append(line.strip())

    # unterminated fence (shouldn't happen in well-formed drafts) -- flush as code
    if fence_lines:
        blocks.append(("code", "\n".join(fence_lines)))
    flush_all()
    return blocks


def is_separator_row(cells):
    return all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip() != "")


def table_rows(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if is_separator_row(cells):
            continue
        rows.append(cells)
    return rows


# ---------- Inline formatting ----------

def add_runs(paragraph, text, base_size=12, base_italic=False, base_bold=False):
    parts = TOKEN_RE.split(text)
    for part in parts:
        if part == "":
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.italic = base_italic
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.bold = base_bold
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.bold = base_bold
            run.italic = base_italic
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic
        run.font.size = Pt(base_size)
        run.font.name = "Times New Roman" if not (part.startswith("`") and part.endswith("`")) else run.font.name


# ---------- Document building ----------

def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.left_margin = Cm(4)
        section.top_margin = Cm(4)
        section.right_margin = Cm(3)
        section.bottom_margin = Cm(3)


def add_title_page(doc):
    def p(text, size=12, bold=False, italic=False, space_after=0, space_before=0):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return para

    for _ in range(3):
        doc.add_paragraph()
    p("ANALISIS OPTIMASI REAL-TIME PIPELINE NVIDIA DEEPSTREAM UNTUK APLIKASI "
      "ADAS BERBASIS EDGE DEVICE", size=16, bold=True, space_after=6)
    p("(Naskah gabungan Bab I sampai IV, draf untuk diperiksa dosen pembimbing)", size=11,
      italic=True, space_after=0)
    for _ in range(4):
        doc.add_paragraph()
    p("SKRIPSI", size=13, bold=True, space_after=0)
    for _ in range(4):
        doc.add_paragraph()
    p("Disusun dan diajukan oleh", size=12, space_after=6)
    p("PERDI", size=13, bold=True, space_after=0)
    p("D121221015", size=12, space_after=0)
    for _ in range(6):
        doc.add_paragraph()
    p("DEPARTEMEN TEKNIK INFORMATIKA", size=12, bold=True, space_after=0)
    p("FAKULTAS TEKNIK", size=12, bold=True, space_after=0)
    p("UNIVERSITAS HASANUDDIN", size=12, bold=True, space_after=0)
    p("2026", size=12, bold=True, space_after=0)
    doc.add_page_break()


def add_toc(doc):
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("DAFTAR ISI")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    note = doc.add_paragraph()
    note_run = note.add_run(
        "(Klik kanan pada area di bawah ini di Microsoft Word, lalu pilih “Update "
        "Field” untuk menampilkan daftar isi otomatis.)"
    )
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_run.font.name = "Times New Roman"

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    doc.add_page_break()


def add_heading(doc, level, text):
    if level == 1:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(18)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        para.style = doc.styles["Heading 1"]
        for r in para.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
            r.font.color.rgb = None
    elif level == 2:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        para.style = doc.styles["Heading 2"]
        for r in para.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
    else:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        para.style = doc.styles["Heading 3"]
        for r in para.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True


def add_paragraph_block(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    add_runs(para, text)


def add_list_block(doc, list_data):
    kind, items = list_data
    style_name = "List Number" if kind == "num" else "List Bullet"
    for item in items:
        para = doc.add_paragraph(style=style_name)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(4)
        add_runs(para, item)


def add_table_block(doc, table_lines):
    rows = table_rows(table_lines)
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx in range(n_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, text, base_size=10.5, base_bold=(r_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc, code_text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    lines = code_text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run().add_break()
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(11)


def _resolve_mermaid_image(mermaid_source):
    for marker, filename, caption in MERMAID_IMAGE_MAP:
        if marker in mermaid_source:
            return os.path.join(DIAGRAMS_DIR, filename), caption
    return None, None


def add_mermaid_block(doc, mermaid_source):
    image_path, caption = _resolve_mermaid_image(mermaid_source)
    if image_path and os.path.exists(image_path):
        add_image_block(doc, image_path, caption)
    else:
        # Fallback: no pre-rendered match found -- keep the raw source visible
        # (as a monospace block) rather than silently dropping the diagram.
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run(
            "[Diagram Mermaid tanpa render gambar yang cocok -- sumber ditampilkan mentah]"
        )
        note_run.italic = True
        note_run.font.name = "Times New Roman"
        note_run.font.size = Pt(10)
        add_code_block(doc, mermaid_source)


def add_image_block(doc, image_path, caption=None, base_dir=None):
    if not os.path.isabs(image_path):
        image_path = os.path.normpath(os.path.join(base_dir or DRAFT_DIR, image_path))
    if not os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[Gambar tidak ditemukan: {image_path}]")
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        return
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(6)
    pic_para.paragraph_format.space_after = Pt(2)
    run = pic_para.add_run()
    run.add_picture(image_path, width=Cm(14))
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(10)
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.name = "Times New Roman"
        cap_run.font.size = Pt(10.5)


def render_blocks(doc, blocks, skip_first_h1=False):
    first_h1_seen = False
    for kind, data in blocks:
        if kind == "h1":
            if skip_first_h1 and not first_h1_seen:
                first_h1_seen = True
                add_heading(doc, 1, data)
                continue
            add_heading(doc, 1, data)
        elif kind == "h2":
            add_heading(doc, 2, data)
        elif kind == "h3":
            add_heading(doc, 3, data)
        elif kind == "para":
            add_paragraph_block(doc, data)
        elif kind == "list":
            add_list_block(doc, data)
        elif kind == "table":
            add_table_block(doc, data)
        elif kind == "code":
            add_code_block(doc, data)
        elif kind == "mermaid":
            add_mermaid_block(doc, data)
        elif kind == "image":
            alt, path = data
            add_image_block(doc, path, caption=alt, base_dir=DRAFT_DIR)


VERIFICATION_NOTES = [
    ("Prioritas tinggi",
     "Penggantian rumusan masalah #3 (sumbu Deep Learning Accelerator/DLA menjadi "
     "perbandingan efisiensi komputasi algoritma tracking NvDCF vs. NvSORT) pada "
     "Bab I belum dikonfirmasi ke dosen pembimbing. Penulis memilih melanjutkan "
     "penulisan draf terlebih dahulu hingga Bab IV selesai, dengan rencana merevisi "
     "pada bimbingan berikutnya apabila perubahan ini tidak disetujui."),
    ("Prioritas tinggi",
     "Verifikasi akurasi as-deployed FP16 pada pembahasan hasil baru selesai pada "
     "tahap infrastruktur kode (probe dump deteksi, konversi video evaluasi, dan "
     "skrip penghitungan mAP). Ekspor dataset validasi ke perangkat Jetson, eksekusi "
     "pengujian yang sesungguhnya, dan pembaruan angka luaran belum dilakukan. Bab IV "
     "sudah mencantumkan keterbatasan ini secara eksplisit, sehingga kesimpulan "
     "terkait trade-off akurasi masih bersyarat pada proxy FP32 dan tidak diklaim "
     "pasti."),
    ("Prioritas sedang",
     "Identifikasi varian Jetson Orin Nano 4GB pada Bab II diturunkan dari pembacaan "
     "mode daya maksimum (nvpmodel -q), bukan dari inspeksi label fisik modul. "
     "Verifikasi independen (mis. lewat pembacaan device-tree perangkat) belum "
     "dilakukan."),
    ("Prioritas sedang",
     "Versi persis GStreamer, GLib, dan CUDA Toolkit yang terpasang pada perangkat "
     "pengujian belum dicatat sebagai metadata reproducibility (mis. lewat "
     "gst-inspect-1.0 --version, nvcc --version, dpkg -l)."),
    ("Prioritas sedang",
     "Hyperparameter pelatihan model pretrained (rasio split, seed, jumlah epoch, "
     "ukuran batch, optimizer, augmentasi, sumber bobot pretrained) belum "
     "dipindahkan dari catatan proses pelatihan ke dokumentasi resmi proyek."),
    ("Prioritas sedang",
     "Ambang selisih mAP FP16-vs-FP32 yang dianggap “dapat diabaikan” belum "
     "dikunci ke angka final, akan ditentukan berdasarkan referensi literatur "
     "setelah data aktual tersedia, untuk menghindari penyesuaian kriteria setelah "
     "melihat hasil. Ambang FPS target real-time sudah dikunci ke 30 FPS."),
    ("Prioritas rendah",
     "Nama laboratorium spesifik tempat pelaksanaan pengujian perangkat keras belum "
     "dikonfirmasi penulis, belum tercatat secara eksplisit di dokumen proyek mana "
     "pun."),
    ("Prioritas rendah",
     "Satu sitasi pada Daftar Pustaka proposal awal, yaitu Wu dkk. (2024, “Road "
     "object detection based on improved YOLOv8 for real-time traffic scenarios”, "
     "Sensors 24(3), 1023), telah dihapus setelah verifikasi DOI melalui Crossref "
     "menunjukkan bahwa rujukan tersebut sebenarnya milik artikel lain yang tidak "
     "berkaitan. Sitasi ini tidak pernah dipakai sebagai rujukan substantif di bagian "
     "mana pun pada naskah ini, namun tetap perlu disampaikan kepada dosen pembimbing "
     "karena merupakan bagian dari Daftar Pustaka yang telah disetujui pada seminar "
     "proposal."),
    ("Prioritas rendah",
     "Tujuh sitasi pada Bab I (Choi, Nigade, Zhang, Ruiz-Barroso, Suder, Seyfipoor, "
     "Shah) masih dirangkum setingkat klaim tematik mengikuti proposal, perlu dibaca "
     "ulang dari sumber aslinya apabila versi final membutuhkan ringkasan "
     "metodologi/hasil yang lebih dalam per jurnal."),
]


def add_verification_appendix(doc):
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("CATATAN VERIFIKASI DAN TINDAK LANJUT PENULIS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    heading.paragraph_format.space_after = Pt(12)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.line_spacing = 1.5
    add_runs(
        intro,
        "Daftar berikut merangkum butir-butir yang masih memerlukan verifikasi, "
        "pelengkapan data, atau konfirmasi pembimbing sebelum naskah gabungan ini "
        "dapat dianggap final. Daftar ini sengaja dipisahkan dari isi Bab I–IV agar "
        "alur akademis pada bab-bab tersebut tetap terbaca sebagai naskah yang sudah "
        "tuntas, sambil tetap menjaga transparansi mengenai bagian yang belum "
        "diselesaikan (sesuai prinsip non-fabrikasi data pada proyek ini). Naskah ini "
        "sudah mencakup keempat bab (Pendahuluan, Metode Penelitian, Hasil dan "
        "Pembahasan, serta Kesimpulan dan Saran) berdasarkan data eksekusi 60 run "
        "benchmark aktual (6 model × 2 tracker × 5 repetisi) pada Jetson Orin Nano; "
        "satu-satunya bagian hasil yang masih tertunda adalah verifikasi akurasi "
        "as-deployed FP16 (lihat butir di bawah).",
    )

    for priority, note in VERIFICATION_NOTES:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(f"[{priority}] ")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        add_runs(para, note)


def main():
    doc = Document()
    set_default_style(doc)
    add_title_page(doc)
    add_toc(doc)

    for idx, path in enumerate(CHAPTER_FILES):
        with open(path, encoding="utf-8") as f:
            text = f.read()
        blocks = parse_blocks(text)
        if idx > 0:
            doc.add_page_break()
        render_blocks(doc, blocks)

    add_verification_appendix(doc)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
